import docx
from docx.shared import Pt, Inches

def main():
    doc = docx.Document('template.docx')
    
    # Add a few empty paragraphs to push down below the letterhead header
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Title
    title = doc.add_paragraph()
    title.alignment = 1 # Center
    run = title.add_run("INVOICE")
    run.font.size = Pt(24)
    run.font.bold = True
    
    doc.add_paragraph()
    doc.add_paragraph("Date: {date}")
    
    doc.add_paragraph()
    doc.add_paragraph("Billed To:")
    p = doc.add_paragraph("{client_name}")
    p.runs[0].font.bold = True
    
    doc.add_paragraph()
    
    # Table
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Description'
    hdr_cells[0].paragraphs[0].runs[0].font.bold = True
    hdr_cells[1].text = 'Amount'
    hdr_cells[1].paragraphs[0].runs[0].font.bold = True
    
    row_cells = table.rows[1].cells
    row_cells[0].text = 'Campaign: {campaign_name}\nBrand: {brand_name}'
    row_cells[1].text = '{amount}'
    
    doc.add_paragraph()
    
    p_total = doc.add_paragraph("Total: {amount}")
    p_total.alignment = 2 # Right
    p_total.runs[0].font.bold = True
    
    doc.save('invoice_template.docx')

if __name__ == '__main__':
    main()
